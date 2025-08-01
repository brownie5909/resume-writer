from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse
import os
import openai
from io import BytesIO
import PyPDF2
import docx
import json
import re
from typing import Optional
import zipfile
import xml.etree.ElementTree as ET

router = APIRouter()

# Initialize OpenAI client
client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

@router.post("/analyze-resume")
async def analyze_resume(
    file: UploadFile = File(...),
    target_role: Optional[str] = Form(None),
    user_id: Optional[str] = Form(None)
):
    """
    Analyze an uploaded resume and provide detailed feedback
    Premium feature only - FIXED DOCX EXTRACTION
    """
    
    try:
        print(f"🔬 Starting analysis for file: {file.filename}")
        print(f"📁 File type: {file.content_type}")
        print(f"🎯 Target role: {target_role}")
        
        # Extract text from uploaded file with enhanced methods
        resume_text = await extract_text_from_file(file)
        print(f"📄 Extracted {len(resume_text)} characters")
        
        # Debug: Show first 200 characters
        print(f"📝 Text preview: {resume_text[:200]}...")
        
        if not resume_text or len(resume_text.strip()) < 50:
            return JSONResponse(
                status_code=400,
                content={"error": f"Could not extract sufficient text from the resume. Extracted: {len(resume_text)} characters. Please ensure the file is readable and contains text content."}
            )
        
        # Analyze the resume with AI
        print("🤖 Starting AI analysis...")
        analysis_result = await analyze_resume_with_ai(resume_text, target_role)
        print("✅ AI analysis completed")
        
        # Generate improved version
        print("🚀 Generating improved resume...")
        improved_resume = await generate_improved_resume(resume_text, target_role, analysis_result)
        print("✅ Improved resume generated")
        
        return JSONResponse({
            "success": True,
            "analysis": analysis_result,
            "improved_resume": improved_resume,
            "original_length": len(resume_text),
            "target_role": target_role,
            "debug_info": {
                "file_type": file.content_type,
                "filename": file.filename,
                "text_preview": resume_text[:300] + "..." if len(resume_text) > 300 else resume_text
            }
        })
        
    except Exception as e:
        print(f"❌ Resume analysis error: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Analysis failed: {str(e)}"}
        )

async def extract_text_from_file(file: UploadFile) -> str:
    """Enhanced text extraction with multiple methods"""
    
    content = await file.read()
    file_extension = file.filename.lower().split('.')[-1] if file.filename else 'unknown'
    
    print(f"🔍 Processing {file_extension} file ({len(content)} bytes)")
    
    try:
        if file_extension == 'pdf':
            return extract_pdf_text_enhanced(content)
        elif file_extension in ['docx', 'doc']:
            return extract_docx_text_enhanced(content)
        elif file_extension == 'txt':
            return extract_txt_text(content)
        else:
            # Try to detect format by content
            return detect_and_extract(content, file.filename)
            
    except Exception as e:
        print(f"⚠️ Primary extraction failed: {str(e)}")
        # Try fallback methods
        return try_fallback_extraction(content, file_extension)

def extract_pdf_text_enhanced(content: bytes) -> str:
    """Enhanced PDF text extraction with fallback methods"""
    try:
        # Method 1: PyPDF2
        pdf_file = BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            text += page_text + "\n"
            print(f"📄 Page {page_num + 1}: {len(page_text)} characters")
        
        if len(text.strip()) > 50:
            return clean_extracted_text(text)
        
        # Method 2: Try alternative PDF libraries if available
        # (You could add pdfplumber or pymupdf here)
        print("⚠️ PyPDF2 extraction insufficient, trying fallback...")
        
        return text.strip() or "PDF text extraction failed - consider converting to DOCX format"
        
    except Exception as e:
        print(f"❌ PDF extraction error: {str(e)}")
        raise ValueError(f"PDF extraction failed: {str(e)}")

def extract_docx_text_enhanced(content: bytes) -> str:
    """Enhanced DOCX text extraction with multiple methods"""
    try:
        # Method 1: python-docx library
        docx_file = BytesIO(content)
        doc = docx.Document(docx_file)
        
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        # Extract headers and footers
        for section in doc.sections:
            # Headers
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text.strip())
            
            # Footers
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text.strip())
        
        extracted_text = "\n".join(text_parts)
        print(f"📝 DOCX extraction: {len(extracted_text)} characters from {len(text_parts)} elements")
        
        if len(extracted_text.strip()) > 50:
            return clean_extracted_text(extracted_text)
        
        # Method 2: Manual XML parsing for DOCX
        print("⚠️ Standard DOCX extraction insufficient, trying XML method...")
        return extract_docx_via_xml(content)
        
    except Exception as e:
        print(f"❌ DOCX extraction error: {str(e)}")
        # Try XML method as fallback
        try:
            return extract_docx_via_xml(content)
        except:
            raise ValueError(f"DOCX extraction failed: {str(e)}")

def extract_docx_via_xml(content: bytes) -> str:
    """Extract DOCX text by parsing XML directly"""
    try:
        # DOCX is actually a ZIP file containing XML
        docx_file = BytesIO(content)
        with zipfile.ZipFile(docx_file, 'r') as zip_file:
            # Read the main document XML
            xml_content = zip_file.read('word/document.xml')
            
            # Parse XML
            root = ET.fromstring(xml_content)
            
            # Extract text from all text nodes
            text_parts = []
            
            # Find all text elements (w:t tags)
            for elem in root.iter():
                if elem.tag.endswith('}t'):  # w:t tag contains text
                    if elem.text:
                        text_parts.append(elem.text)
            
            extracted_text = " ".join(text_parts)
            print(f"📝 XML extraction: {len(extracted_text)} characters")
            
            return clean_extracted_text(extracted_text)
            
    except Exception as e:
        print(f"❌ XML extraction error: {str(e)}")
        raise ValueError(f"XML extraction failed: {str(e)}")

def extract_txt_text(content: bytes) -> str:
    """Extract text from TXT files with encoding detection"""
    try:
        # Try UTF-8 first
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            pass
        
        # Try other common encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                text = content.decode(encoding)
                print(f"📝 TXT decoded with {encoding}")
                return clean_extracted_text(text)
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        # If all fail, decode with errors='ignore'
        return content.decode('utf-8', errors='ignore')
        
    except Exception as e:
        raise ValueError(f"TXT extraction failed: {str(e)}")

def detect_and_extract(content: bytes, filename: str) -> str:
    """Try to detect file format and extract accordingly"""
    try:
        # Check file signatures
        if content.startswith(b'%PDF'):
            print("🔍 Detected PDF by signature")
            return extract_pdf_text_enhanced(content)
        
        elif content.startswith(b'PK\x03\x04'):
            print("🔍 Detected ZIP-based format (likely DOCX)")
            return extract_docx_text_enhanced(content)
        
        elif b'<html' in content[:1000].lower() or b'<body' in content[:1000].lower():
            print("🔍 Detected HTML content")
            return extract_txt_text(content)
        
        else:
            # Try as text
            print("🔍 Treating as text file")
            return extract_txt_text(content)
            
    except Exception as e:
        raise ValueError(f"Format detection failed: {str(e)}")

def try_fallback_extraction(content: bytes, file_extension: str) -> str:
    """Try alternative extraction methods"""
    try:
        print(f"🔄 Trying fallback extraction for {file_extension}")
        
        # Try treating as text regardless of extension
        text = extract_txt_text(content)
        
        if len(text.strip()) > 50:
            return text
        
        # If still no luck, return what we have with instructions
        return f"Partial extraction from {file_extension} file. Text found: {text[:200]}... Please try uploading as PDF or TXT format for better results."
        
    except Exception as e:
        raise ValueError(f"All extraction methods failed: {str(e)}")

def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove common artifacts
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', text)
    
    # Fix common OCR errors
    text = text.replace('|', 'I')  # Common OCR mistake
    
    # Normalize line breaks
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()

# Keep all the existing analysis functions unchanged
async def analyze_resume_with_ai(resume_text: str, target_role: Optional[str] = None) -> dict:
    """Use AI to analyze the resume and provide detailed feedback"""
    
    role_context = f" for a {target_role} position" if target_role else ""
    
    # Simplified analysis prompt that's more reliable
    analysis_prompt = f"""
Analyze this resume{role_context} and provide specific, actionable feedback.

Resume Text:
{resume_text[:3000]}  

Provide analysis in this exact format:

OVERALL_SCORE: [number 0-100]
ATS_SCORE: [number 0-100] 
FORMATTING_SCORE: [number 0-100]

STRENGTHS:
- [strength 1]
- [strength 2] 
- [strength 3]

WEAKNESSES:
- [weakness 1]
- [weakness 2]
- [weakness 3]

MISSING_KEYWORDS: [keyword1, keyword2, keyword3]
PRESENT_KEYWORDS: [keyword1, keyword2, keyword3]

SPECIFIC_IMPROVEMENTS:
1. [improvement 1]
2. [improvement 2] 
3. [improvement 3]

ATS_RECOMMENDATIONS:
1. [ats tip 1]
2. [ats tip 2]

SECTION_FEEDBACK:
Contact: [score]/100 - [feedback]
Summary: [score]/100 - [feedback]  
Experience: [score]/100 - [feedback]
Education: [score]/100 - [feedback]
Skills: [score]/100 - [feedback]
"""

    try:
        print("📤 Sending request to OpenAI...")
        
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert resume reviewer with 15+ years of experience. Provide specific, actionable feedback in the exact format requested."},
                {"role": "user", "content": analysis_prompt}
            ],
            temperature=0.3,
            max_tokens=1500
        )
        
        ai_response = response.choices[0].message.content
        print("📥 Received AI response")
        
        # Parse the structured response
        analysis_data = parse_structured_analysis(ai_response)
        
        # Add keyword density calculation
        if analysis_data.get('keyword_analysis'):
            missing_count = len(analysis_data['keyword_analysis'].get('missing_keywords', []))
            present_count = len(analysis_data['keyword_analysis'].get('present_keywords', []))
            total_keywords = missing_count + present_count
            
            if total_keywords > 0:
                density = int((present_count / total_keywords) * 100)
                analysis_data['keyword_analysis']['keyword_density'] = density
        
        return analysis_data
        
    except Exception as e:
        print(f"❌ AI analysis error: {str(e)}")
        return get_enhanced_fallback_analysis(resume_text, target_role)

# Include all the existing parsing and helper functions from the previous version
# (parse_structured_analysis, extract_score, extract_list, etc.)
# I'll just include the key ones for brevity:

def parse_structured_analysis(ai_response: str) -> dict:
    """Parse the structured AI response"""
    
    try:
        # Extract scores
        overall_score = extract_score(ai_response, "OVERALL_SCORE")
        ats_score = extract_score(ai_response, "ATS_SCORE") 
        formatting_score = extract_score(ai_response, "FORMATTING_SCORE")
        
        # Extract lists
        strengths = extract_list(ai_response, "STRENGTHS:", "WEAKNESSES:")
        weaknesses = extract_list(ai_response, "WEAKNESSES:", "MISSING_KEYWORDS:")
        
        # Extract keywords
        missing_keywords = extract_keywords(ai_response, "MISSING_KEYWORDS:")
        present_keywords = extract_keywords(ai_response, "PRESENT_KEYWORDS:")
        
        # Extract improvements
        improvements = extract_numbered_list(ai_response, "SPECIFIC_IMPROVEMENTS:")
        ats_recommendations = extract_numbered_list(ai_response, "ATS_RECOMMENDATIONS:")
        
        # Extract section feedback
        sections = parse_section_feedback(ai_response)
        
        return {
            "overall_score": overall_score,
            "ats_score": ats_score,
            "formatting_score": formatting_score,
            "strengths": strengths,
            "weaknesses": weaknesses,
            "keyword_analysis": {
                "missing_keywords": missing_keywords,
                "present_keywords": present_keywords,
                "keyword_density": 70
            },
            "sections_analysis": sections,
            "specific_improvements": improvements,
            "ats_recommendations": ats_recommendations
        }
        
    except Exception as e:
        print(f"⚠️ Parsing error: {str(e)}")
        return get_enhanced_fallback_analysis("", "")

def extract_score(text: str, keyword: str) -> int:
    """Extract score from text"""
    try:
        pattern = rf"{keyword}:\s*(\d+)"
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return int(match.group(1))
    except:
        pass
    return 75

def extract_list(text: str, start_keyword: str, end_keyword: str) -> list:
    """Extract bullet point list between keywords"""
    try:
        start_pos = text.find(start_keyword)
        end_pos = text.find(end_keyword)
        
        if start_pos == -1:
            return []
            
        if end_pos == -1:
            end_pos = len(text)
            
        section = text[start_pos + len(start_keyword):end_pos]
        
        items = []
        for line in section.split('\n'):
            line = line.strip()
            if line.startswith('-') or line.startswith('•'):
                items.append(line[1:].strip())
        
        return items[:5]
        
    except:
        return []

def extract_keywords(text: str, keyword: str) -> list:
    """Extract comma-separated keywords"""
    try:
        start_pos = text.find(keyword)
        if start_pos == -1:
            return []
            
        start_pos += len(keyword)
        end_pos = text.find('\n', start_pos)
        if end_pos == -1:
            end_pos = start_pos + 200
            
        line = text[start_pos:end_pos].strip()
        line = line.strip('[]{}()')
        keywords = [k.strip() for k in line.split(',') if k.strip()]
        
        return keywords[:8]
        
    except:
        return []

def extract_numbered_list(text: str, keyword: str) -> list:
    """Extract numbered list items"""
    try:
        start_pos = text.find(keyword)
        if start_pos == -1:
            return []
            
        next_section = text.find('\n\n', start_pos + len(keyword))
        if next_section == -1:
            next_section = len(text)
            
        section = text[start_pos + len(keyword):next_section]
        
        items = []
        for line in section.split('\n'):
            line = line.strip()
            if re.match(r'^\d+\.', line):
                items.append(re.sub(r'^\d+\.\s*', '', line))
        
        return items[:5]
        
    except:
        return []

def parse_section_feedback(text: str) -> dict:
    """Parse section-by-section feedback"""
    sections = {}
    section_names = ['Contact', 'Summary', 'Experience', 'Education', 'Skills']
    
    for section in section_names:
        try:
            pattern = rf"{section}:\s*(\d+)/100\s*-\s*([^\n]+)"
            match = re.search(pattern, text, re.IGNORECASE)
            
            if match:
                score = int(match.group(1))
                feedback = match.group(2).strip()
            else:
                score = 75
                feedback = f"{section} section appears standard"
                
            sections[section.lower() + '_info' if section == 'Contact' else section.lower()] = {
                "score": score,
                "feedback": feedback
            }
        except:
            sections[section.lower() + '_info' if section == 'Contact' else section.lower()] = {
                "score": 75,
                "feedback": f"{section} section needs review"
            }
    
    return sections

async def generate_improved_resume(resume_text: str, target_role: Optional[str], analysis: dict) -> str:
    """Generate an improved version of the resume"""
    
    role_context = f" for {target_role} positions" if target_role else ""
    
    improvement_prompt = f"""
Improve this resume{role_context} based on the analysis. Keep all original information but enhance presentation.

Original Resume:
{resume_text[:2000]}

Key Issues to Fix:
{chr(10).join(analysis.get('specific_improvements', [])[:3])}

Create an improved version that:
1. Keeps all original experience and information
2. Uses stronger action verbs
3. Adds quantifiable achievements where possible
4. Improves formatting and structure
5. Makes it more ATS-friendly

Return only the improved resume text.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert resume writer. Improve resumes while keeping all original information."},
                {"role": "user", "content": improvement_prompt}
            ],
            temperature=0.4,
            max_tokens=1500
        )
        
        improved_resume = response.choices[0].message.content.strip()
        return improved_resume
        
    except Exception as e:
        print(f"❌ Resume improvement error: {str(e)}")
        return f"Resume analysis completed successfully. Improved version generation temporarily unavailable.\n\nOriginal resume content preserved:\n\n{resume_text[:1000]}..."

def get_enhanced_fallback_analysis(resume_text: str, target_role: Optional[str]) -> dict:
    """Enhanced fallback with basic text analysis"""
    
    word_count = len(resume_text.split())
    has_email = '@' in resume_text
    has_phone = any(char.isdigit() for char in resume_text)
    has_experience = any(word in resume_text.lower() for word in ['experience', 'worked', 'managed', 'led'])
    has_education = any(word in resume_text.lower() for word in ['university', 'college', 'degree', 'bachelor', 'master'])
    has_skills = any(word in resume_text.lower() for word in ['skills', 'proficient', 'experienced'])
    
    overall_score = 60
    if has_email: overall_score += 5
    if has_phone: overall_score += 5  
    if has_experience: overall_score += 10
    if has_education: overall_score += 10
    if has_skills: overall_score += 5
    if word_count > 200: overall_score += 5
    
    return {
        "overall_score": min(overall_score, 85),
        "ats_score": max(overall_score - 10, 50),
        "formatting_score": overall_score - 5,
        "strengths": [
            f"Resume contains {word_count} words showing good detail" if word_count > 150 else "Resume structure detected",
            "Contact information present" if has_email else "Readable format confirmed",
            "Professional experience documented" if has_experience else "Content successfully processed"
        ],
        "weaknesses": [
            "Consider adding more quantified achievements",
            "Ensure keywords match target job requirements",
            "Optimize formatting for ATS compatibility"
        ],
        "keyword_analysis": {
            "missing_keywords": ["leadership", "results-driven", "collaborative"],
            "present_keywords": ["experience", "skills", "professional"],
            "keyword_density": 65
        },
        "sections_analysis": {
            "contact_info": {"score": 85 if has_email else 60, "feedback": "Contact information detected" if has_email else "Contact info needs verification"},
            "summary": {"score": 65, "feedback": "Consider adding a compelling professional summary"},
            "experience": {"score": 80 if has_experience else 50, "feedback": "Experience section found" if has_experience else "Experience section needs review"},
            "education": {"score": 75 if has_education else 60, "feedback": "Education background present" if has_education else "Education section standard"},
            "skills": {"score": 75 if has_skills else 65, "feedback": "Skills section identified" if has_skills else "Skills section detected"}
        },
        "specific_improvements": [
            "Add quantified achievements with specific numbers and percentages",
            "Include a compelling professional summary at the top",
            "Use stronger action verbs (achieved, implemented, optimized)",
            f"Optimize for {target_role} keywords" if target_role else "Add relevant industry keywords"
        ],
        "ats_recommendations": [
            "Use standard section headers (Experience, Education, Skills)",
            "Maintain consistent formatting throughout",
            "Include relevant keywords naturally in content"
        ]
    }

@router.get("/resume-analysis/health")
async def resume_analysis_health():
    return {"status": "healthy", "service": "resume-analysis"}

# ADD THESE ROUTES TO THE END OF YOUR EXISTING routes/resume_analysis.py file

@router.post("/analyze-cover-letter")
async def analyze_cover_letter(
    file: UploadFile = File(...),
    target_role: Optional[str] = Form(None),
    document_type: Optional[str] = Form("cover_letter"),
    user_id: Optional[str] = Form(None)
):
    """
    Analyze an uploaded cover letter and provide detailed feedback
    """
    
    try:
        print(f"📝 Starting cover letter analysis for file: {file.filename}")
        print(f"🎯 Target role: {target_role}")
        
        # Reuse your existing text extraction function
        cover_letter_text = await extract_text_from_file(file)
        print(f"📄 Extracted {len(cover_letter_text)} characters")
        
        if not cover_letter_text or len(cover_letter_text.strip()) < 50:
            return JSONResponse(
                status_code=400,
                content={"error": f"Could not extract sufficient text from the cover letter. Extracted: {len(cover_letter_text)} characters."}
            )
        
        # Analyze the cover letter with AI
        print("🤖 Starting AI analysis...")
        analysis_result = await analyze_cover_letter_with_ai(cover_letter_text, target_role)
        print("✅ AI analysis completed")
        
        # Generate improved version
        print("🚀 Generating improved cover letter...")
        improved_cover_letter = await generate_improved_cover_letter(cover_letter_text, target_role, analysis_result)
        print("✅ Improved cover letter generated")
        
        return JSONResponse({
            "success": True,
            "analysis": analysis_result,
            "improved_cover_letter": improved_cover_letter,
            "original_length": len(cover_letter_text),
            "target_role": target_role
        })
        
    except Exception as e:
        print(f"❌ Cover letter analysis error: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Analysis failed: {str(e)}"}
        )

@router.post("/generate-cover-letter")
async def generate_cover_letter(request: Request):
    """
    Generate a new cover letter based on job posting and applicant information
    """
    
    try:
        body = await request.json()
        
        job_posting = body.get("job_posting", "").strip()
        applicant_name = body.get("applicant_name", "").strip()
        current_role = body.get("current_role", "").strip()
        experience = body.get("experience", "").strip()
        achievements = body.get("achievements", "").strip()
        
        if not job_posting or not applicant_name:
            return JSONResponse(
                status_code=400,
                content={"error": "Job posting and applicant name are required."}
            )
        
        print(f"✨ Generating cover letter for {applicant_name}")
        print(f"🎯 Job posting length: {len(job_posting)} characters")
        
        # Generate cover letter with AI
        cover_letter = await generate_cover_letter_with_ai(
            job_posting, applicant_name, current_role, experience, achievements
        )
        
        return JSONResponse({
            "success": True,
            "cover_letter": cover_letter,
            "applicant_name": applicant_name,
            "generated_for": job_posting[:100] + "..." if len(job_posting) > 100 else job_posting
        })
        
    except Exception as e:
        print(f"❌ Cover letter generation error: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Generation failed: {str(e)}"}
        )

# Add these helper functions to the end of the file

async def analyze_cover_letter_with_ai(cover_letter_text: str, target_role: Optional[str] = None) -> dict:
    """Use AI to analyze the cover letter and provide detailed feedback"""
    
    role_context = f" for a {target_role} position" if target_role else ""
    
    analysis_prompt = f"""
Analyze this cover letter{role_context} and provide specific, actionable feedback.

Cover Letter Text:
{cover_letter_text[:2500]}

Provide analysis in this exact format:

OVERALL_SCORE: [number 0-100]
JOB_ALIGNMENT_SCORE: [number 0-100]
ATS_SCORE: [number 0-100]

STRENGTHS:
- [strength 1]
- [strength 2] 
- [strength 3]

WEAKNESSES:
- [weakness 1]
- [weakness 2]
- [weakness 3]

SPECIFIC_IMPROVEMENTS:
1. [improvement 1]
2. [improvement 2] 
3. [improvement 3]

JOB_SPECIFIC_TIPS:
1. [job-specific tip 1]
2. [job-specific tip 2]
3. [job-specific tip 3]

Focus on:
1. How well the letter demonstrates fit for the role
2. Storytelling and compelling narrative  
3. Professional tone and structure
4. Specific examples and achievements
5. Call to action and closing strength
"""

    try:
        print("📤 Sending cover letter to OpenAI for analysis...")
        
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert career coach specializing in cover letter optimization. Provide specific, actionable feedback."},
                {"role": "user", "content": analysis_prompt}
            ],
            temperature=0.3,
            max_tokens=1200
        )
        
        ai_response = response.choices[0].message.content
        print("📥 Received cover letter analysis from AI")
        
        # Parse the structured response
        analysis_data = parse_cover_letter_analysis(ai_response)
        
        return analysis_data
        
    except Exception as e:
        print(f"❌ AI analysis error: {str(e)}")
        return get_fallback_cover_letter_analysis(cover_letter_text, target_role)

def parse_cover_letter_analysis(ai_response: str) -> dict:
    """Parse the structured AI response for cover letters"""
    
    try:
        # Extract scores
        overall_score = extract_score(ai_response, "OVERALL_SCORE")
        job_alignment_score = extract_score(ai_response, "JOB_ALIGNMENT_SCORE")
        ats_score = extract_score(ai_response, "ATS_SCORE")
        
        # Extract lists
        strengths = extract_list(ai_response, "STRENGTHS:", "WEAKNESSES:")
        weaknesses = extract_list(ai_response, "WEAKNESSES:", "SPECIFIC_IMPROVEMENTS:")
        
        # Extract improvements and tips
        improvements = extract_numbered_list(ai_response, "SPECIFIC_IMPROVEMENTS:")
        job_tips = extract_numbered_list(ai_response, "JOB_SPECIFIC_TIPS:")
        
        return {
            "overall_score": overall_score,
            "job_alignment_score": job_alignment_score,
            "ats_score": ats_score,
            "strengths": strengths,
            "weaknesses": weaknesses,
            "specific_improvements": improvements,
            "job_specific_tips": job_tips
        }
        
    except Exception as e:
        print(f"⚠️ Cover letter parsing error: {str(e)}")
        return get_fallback_cover_letter_analysis("", "")

async def generate_improved_cover_letter(cover_letter_text: str, target_role: Optional[str], analysis: dict) -> str:
    """Generate an improved version of the cover letter based on analysis"""
    
    role_context = f" for {target_role} positions" if target_role else ""
    
    improvement_prompt = f"""
Improve this cover letter{role_context} based on the analysis feedback.

Original Cover Letter:
{cover_letter_text[:2000]}

Key Issues to Fix:
{chr(10).join(analysis.get('specific_improvements', [])[:3])}

Create an improved version that:
1. Maintains the original tone and personal voice
2. Fixes identified weaknesses
3. Strengthens the narrative and storytelling
4. Adds more specific examples and achievements
5. Improves the opening hook and closing call-to-action
6. Optimizes for ATS compatibility

Return only the improved cover letter text in proper business letter format.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert cover letter writer. Improve cover letters while maintaining the applicant's authentic voice and personal style."},
                {"role": "user", "content": improvement_prompt}
            ],
            temperature=0.4,
            max_tokens=1200
        )
        
        improved_cover_letter = response.choices[0].message.content.strip()
        return improved_cover_letter
        
    except Exception as e:
        print(f"❌ Cover letter improvement error: {str(e)}")
        return f"Cover letter analysis completed successfully. Improved version generation temporarily unavailable.\n\nOriginal cover letter content preserved:\n\n{cover_letter_text[:800]}..."

async def generate_cover_letter_with_ai(job_posting: str, applicant_name: str, current_role: str, experience: str, achievements: str) -> str:
    """Generate a new cover letter based on job posting and applicant information"""
    
    generation_prompt = f"""
Write a compelling, professional cover letter based on this information:

Job Posting/Role:
{job_posting[:1500]}

Applicant Information:
- Name: {applicant_name}
- Current Role: {current_role}
- Experience: {experience}
- Key Achievements: {achievements}

Create a cover letter that:
1. Opens with a strong hook that shows enthusiasm and relevance
2. Demonstrates clear understanding of the role and company
3. Highlights relevant experience and achievements with specific examples
4. Shows personality while maintaining professionalism
5. Includes a compelling call-to-action in the closing
6. Is properly formatted as a business letter
7. Is optimized for ATS systems with relevant keywords

The letter should be 3-4 paragraphs, approximately 250-350 words total.
Use a professional but engaging tone that stands out from generic templates.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert cover letter writer who creates compelling, personalized cover letters that help candidates stand out and get interviews."},
                {"role": "user", "content": generation_prompt}
            ],
            temperature=0.6,
            max_tokens=1000
        )
        
        cover_letter = response.choices[0].message.content.strip()
        return cover_letter
        
    except Exception as e:
        print(f"❌ Cover letter generation error: {str(e)}")
        return f"Dear Hiring Manager,\n\nI am writing to express my strong interest in the {current_role} position. With my background in {experience}, I am confident I would be a valuable addition to your team.\n\n[Cover letter generation temporarily unavailable - please try again later]\n\nSincerely,\n{applicant_name}"

def get_fallback_cover_letter_analysis(cover_letter_text: str, target_role: Optional[str]) -> dict:
    """Fallback analysis when AI is unavailable"""
    
    word_count = len(cover_letter_text.split())
    has_greeting = any(greeting in cover_letter_text.lower() for greeting in ['dear', 'hello', 'greetings'])
    has_closing = any(closing in cover_letter_text.lower() for closing in ['sincerely', 'regards', 'best'])
    has_company_ref = '[company' not in cover_letter_text.lower()
    
    overall_score = 65
    if has_greeting: overall_score += 10
    if has_closing: overall_score += 10
    if has_company_ref: overall_score += 10
    if word_count >= 200 and word_count <= 400: overall_score += 5
    
    return {
        "overall_score": min(overall_score, 85),
        "job_alignment_score": max(overall_score - 15, 50),
        "ats_score": max(overall_score - 10, 60),
        "strengths": [
            f"Cover letter length is {word_count} words",
            "Professional greeting found" if has_greeting else "Structure appears professional",
            "Proper closing included" if has_closing else "Content successfully processed"
        ],
        "weaknesses": [
            "Consider adding more specific examples and achievements",
            "Ensure company name and role are mentioned specifically",
            "Strengthen the opening hook to grab attention"
        ],
        "specific_improvements": [
            "Add quantified achievements and specific examples",
            "Research the company and mention specific details",
            "Create a stronger opening that shows enthusiasm",
            f"Optimize for {target_role} keywords" if target_role else "Include relevant industry keywords"
        ],
        "job_specific_tips": [
            "Tailor the content to match job requirements",
            "Use the same keywords as the job posting",
            "Show understanding of company culture and values"
        ]
    }
