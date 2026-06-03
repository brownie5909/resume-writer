import os
import re
import zipfile
import xml.etree.ElementTree as ET
from io import BytesIO

import PyPDF2
from docx import Document


async def extract_text_from_file(file):
    """
    Extract text from uploaded resume files.
    Supports PDF, DOCX, TXT, and RTF.
    Legacy DOC files are not reliably supported without conversion to DOCX.
    """

    filename = file.filename.lower()

    try:
        content = await file.read()

        if filename.endswith(".pdf"):
            return clean_extracted_text(extract_pdf_text(content))

        if filename.endswith(".docx"):
            return clean_extracted_text(extract_docx_text(content))

        if filename.endswith(".txt"):
            return clean_extracted_text(content.decode("utf-8", errors="ignore"))

        if filename.endswith(".rtf"):
            return clean_extracted_text(extract_rtf_text(content))

        if filename.endswith(".doc"):
            raise ValueError(
                "Legacy .doc files are not supported reliably. Please open the file in Word or Google Docs and save/export it as .docx or PDF, then upload again."
            )

        raise ValueError("Unsupported file format")

    except Exception as e:
        raise Exception(f"File parsing failed: {str(e)}")


def clean_extracted_text(text: str) -> str:
    """Normalise extracted resume text without removing useful line breaks."""
    if not text:
        return ""

    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf_text(content):
    """Extract text from PDF using PyPDF2."""
    text_parts = []

    pdf_reader = PyPDF2.PdfReader(BytesIO(content))

    for page in pdf_reader.pages:
        extracted = page.extract_text()
        if extracted:
            text_parts.append(extracted)

    return "\n".join(text_parts).strip()


def extract_docx_text(content):
    """Extract text from DOCX, including paragraphs, tables, headers/footers and XML fallback."""
    text_parts = []

    doc = Document(BytesIO(content))

    for paragraph in doc.paragraphs:
        if paragraph.text and paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = clean_extracted_text(cell.text)
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(" | ".join(row_text))

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        for paragraph in section.footer.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

    combined_text = clean_extracted_text("\n".join(text_parts))

    if len(combined_text) >= 50:
        return combined_text

    xml_fallback_text = extract_docx_xml_text(content)
    if len(xml_fallback_text) > len(combined_text):
        return xml_fallback_text

    return combined_text


def extract_docx_xml_text(content):
    """Fallback DOCX extraction by reading Word XML directly."""
    text_parts = []
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    xml_files = [
        "word/document.xml",
        "word/header1.xml",
        "word/header2.xml",
        "word/header3.xml",
        "word/footer1.xml",
        "word/footer2.xml",
        "word/footer3.xml",
    ]

    try:
        with zipfile.ZipFile(BytesIO(content)) as docx_zip:
            for xml_file in xml_files:
                if xml_file not in docx_zip.namelist():
                    continue

                xml_content = docx_zip.read(xml_file)
                root = ET.fromstring(xml_content)

                for text_node in root.findall(".//w:t", namespace):
                    if text_node.text and text_node.text.strip():
                        text_parts.append(text_node.text.strip())
    except Exception:
        return ""

    return clean_extracted_text("\n".join(text_parts))


def extract_rtf_text(content):
    """Basic RTF text extraction fallback."""
    raw = content.decode("utf-8", errors="ignore")
    raw = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
    raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", raw)
    raw = raw.replace("{", " ").replace("}", " ")
    raw = raw.replace("\\", " ")
    return raw
